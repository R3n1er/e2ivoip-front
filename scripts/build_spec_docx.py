#!/usr/bin/env python3
"""
Génère SPEC_STRATEGIE_VENTE_MARKETING.docx avec charte graphique E2I VoIP.

Charte appliquée :
- Rouge Principal #E53E3E  → H1, accents, CTA
- Bleu Marine    #2D3848  → H2, H3, en-têtes de tableaux
- Gris Foncé     #1F2937  → texte courant, H4
- Gris Secondaire #818096 → légendes, citations
- Blanc          #FFFFFF  → texte sur fond foncé

Police : Calibri (corps) / Calibri (titres) — équivalent moderne accessible.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from pathlib import Path


# ---------- Charte graphique E2I VoIP ----------
COLOR_RED        = RGBColor(0xE5, 0x3E, 0x3E)
COLOR_MARINE     = RGBColor(0x2D, 0x38, 0x48)
COLOR_DARK       = RGBColor(0x1F, 0x29, 0x37)
COLOR_SECONDARY  = RGBColor(0x81, 0x80, 0x96)
COLOR_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_LIGHT_BG   = "F4F4F6"   # gris très clair pour bandes alternées (hex sans #)
COLOR_RED_BG     = "E53E3E"
COLOR_MARINE_BG  = "2D3848"


# ---------- Helpers de styling bas niveau (XML) ----------
def set_cell_shading(cell, hex_color):
    """Applique une couleur de fond à une cellule de tableau."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="2D3848", size="4"):
    """Bordures fines pour cellules de tableau."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_horizontal_rule(paragraph, color_rgb=COLOR_RED, thickness_pt=2):
    """Ajoute une bordure inférieure (= ligne horizontale) à un paragraphe."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness_pt * 8))   # eighths-of-a-point
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def style_run(run, *, color=COLOR_DARK, bold=False, size=11, font="Calibri", italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


# ---------- Styles document ----------
def configure_default_styles(doc):
    """Configure les styles natifs Word selon la charte."""
    styles = doc.styles

    # Normal
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    # Title (couverture)
    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(32)
    title.font.color.rgb = COLOR_RED
    title.font.bold = True

    # Heading 1
    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.color.rgb = COLOR_RED
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(20)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(16)
    h2.font.color.rgb = COLOR_MARINE
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(13)
    h3.font.color.rgb = COLOR_MARINE
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    # Heading 4
    h4 = styles["Heading 4"]
    h4.font.name = "Calibri"
    h4.font.size = Pt(11)
    h4.font.color.rgb = COLOR_DARK
    h4.font.bold = True
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(2)


# ---------- Builders structure ----------
def add_cover(doc):
    """Page de couverture avec la charte."""
    # Titre principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("E2I VoIP")
    style_run(r, color=COLOR_MARINE, bold=True, size=42)

    # Tagline
    p = doc.add_paragraph()
    r = p.add_run("La téléphonie d'entreprise, simple et évolutive.")
    style_run(r, color=COLOR_SECONDARY, italic=True, size=14)

    # Barre rouge (séparateur)
    sep = doc.add_paragraph()
    add_horizontal_rule(sep, color_rgb=COLOR_RED, thickness_pt=3)

    # Titre du document
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("Spécification — Stratégie de Vente & Marketing")
    style_run(r, color=COLOR_RED, bold=True, size=26)

    # Sous-titre
    p = doc.add_paragraph()
    r = p.add_run("Document de référence interne")
    style_run(r, color=COLOR_DARK, size=14)

    # Bloc métadonnées
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    for label, value in [
        ("Version",  "1.0"),
        ("Date",     "28 avril 2026"),
        ("Auteur",   "Alban Renier — Direction"),
        ("Statut",   "Document de référence — Diffusion interne"),
    ]:
        r = p.add_run(f"{label} : ")
        style_run(r, color=COLOR_MARINE, bold=True, size=11)
        r = p.add_run(f"{value}\n")
        style_run(r, color=COLOR_DARK, size=11)

    # Sources consolidées (encadré)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("Sources consolidées")
    style_run(r, color=COLOR_MARINE, bold=True, size=12)

    sources = [
        "docs/BrandBrief_e2ivoip.md — positionnement, ICP, offres, processus",
        ".planning/PROJECT.md — objectifs business, KPIs, contraintes",
        ".planning/ROADMAP.md — phases d'activation digitale",
        ".planning/REQUIREMENTS.md — exigences mesurables",
        "Travaux GSD (Get Shit Done) — Phases 1, 2, 3",
        "Travaux Superpowers — brainstorming, plans d'exécution",
    ]
    for s in sources:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(s)
        style_run(r, color=COLOR_DARK, size=10)

    doc.add_page_break()


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    add_horizontal_rule(p, color_rgb=COLOR_RED, thickness_pt=2)


def add_h2(doc, text):
    doc.add_paragraph(text, style="Heading 2")


def add_h3(doc, text):
    doc.add_paragraph(text, style="Heading 3")


def add_para(doc, text, *, color=COLOR_DARK, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, color=color, bold=bold, italic=italic, size=size)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        style_run(r, color=COLOR_DARK, size=11)


def add_quote(doc, text):
    """Encadré citation avec accent rouge."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    # bordure gauche rouge
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "E53E3E")
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r = p.add_run(text)
    style_run(r, color=COLOR_MARINE, italic=True, size=12, bold=True)


def add_table(doc, headers, rows, *, header_bg=COLOR_MARINE_BG, alt_bg=COLOR_LIGHT_BG):
    """Tableau stylé charte E2I."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # En-tête
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        style_run(r, color=COLOR_WHITE, bold=True, size=10.5)
        set_cell_shading(cell, header_bg)
        set_cell_borders(cell, color="2D3848")

    # Lignes
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            style_run(r, color=COLOR_DARK, size=10.5)
            if ri % 2 == 1:
                set_cell_shading(cell, alt_bg)
            set_cell_borders(cell, color="D6D6DC")

    # spacing
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_kv_block(doc, pairs):
    """Bloc 'label : valeur' compact."""
    for label, value in pairs:
        p = doc.add_paragraph()
        r = p.add_run(f"{label} : ")
        style_run(r, color=COLOR_MARINE, bold=True, size=11)
        r = p.add_run(value)
        style_run(r, color=COLOR_DARK, size=11)


# ---------- Construction du document ----------
def build_document(output_path):
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    configure_default_styles(doc)

    # === COUVERTURE ===
    add_cover(doc)

    # === 1. SYNTHÈSE EXÉCUTIVE ===
    add_h1(doc, "1. Synthèse Exécutive")
    add_para(doc,
        "E2I VoIP est un opérateur de services télécom et intégrateur VoIP spécialisé dans "
        "l'accompagnement des PME et groupes multisites en téléphonie IP, avec une double "
        "implantation Antilles-Guyane / La Réunion / Métropole et une expertise certifiée 3CX et Yeastar."
    )
    add_para(doc, "La stratégie commerciale s'articule autour de trois leviers :", bold=True)
    add_bullets(doc, [
        "Différenciation par le Customer Success — un Customer Success Manager dédié à chaque client, plutôt qu'une logique transactionnelle.",
        "Proximité géographique — équipes locales dans les DROM et en Métropole, support technique en interne.",
        "Profondeur technique — certification 3CX Silver, maîtrise du Call Flow Designer (CFD), intégrations CRM avancées.",
    ])
    add_para(doc,
        "L'objectif business à court terme (milestone Ship & Measure) est de doubler les leads "
        "qualifiés générés par le site web (baseline actuelle < 20 leads/mois) via :"
    )
    add_bullets(doc, [
        "Le déploiement du redesign Monolithe 2026",
        "L'ajout de social proof crédible (témoignages réels + badge 3CX)",
        "L'optimisation des chemins de conversion (1 CTA primaire above-the-fold par page)",
        "L'instrumentation analytics PostHog pour mesurer avant/après",
    ])

    # === 2. POSITIONNEMENT ===
    add_h1(doc, "2. Positionnement & Promesse de Marque")

    add_h2(doc, "2.1. Slogan officiel")
    add_quote(doc, "« La téléphonie d'entreprise, simple et évolutive. »")

    add_h2(doc, "2.2. Promesse de valeur")
    add_quote(doc,
        "Chaque visiteur doit comprendre en 10 secondes qu'E2I VoIP est un expert local de "
        "confiance en téléphonie d'entreprise, et avoir un chemin clair pour prendre contact."
    )

    add_h2(doc, "2.3. Thèse compétitive")
    add_para(doc,
        "La proximité combinée à l'expertise locale bat les solutions universelles (Teams, "
        "Zoom Phone) sur le segment PME, particulièrement dans les DROM où la maîtrise des "
        "contraintes techniques territoriales est un avantage décisif."
    )

    add_h2(doc, "2.4. Différenciateur central")
    add_para(doc,
        "Là où la majorité des concurrents considèrent la signature du contrat comme un "
        "aboutissement, E2I VoIP la considère comme le début d'une collaboration. Le Customer "
        "Success Manager n'est pas un contact administratif : c'est un partenaire opérationnel "
        "engagé dans la réussite du projet client."
    )

    add_h2(doc, "2.5. Tagline alternative B2B")
    add_quote(doc,
        "Votre opérateur télécom DOM & Métropole. Certifié 3CX. Hébergé en France."
    )

    # === 3. ICP ===
    add_h1(doc, "3. Profil Client Idéal (ICP)")

    add_h2(doc, "3.1. Critères de qualification")
    add_table(doc,
        headers=["Critère", "Cible"],
        rows=[
            ["Taille",       "TPE (≥ 3 utilisateurs) à PME / groupes (60 à 90+ utilisateurs)"],
            ["Géographie",   "Guadeloupe, Martinique, Guyane, La Réunion, France métropolitaine"],
            ["Maturité",     "5+ ans d'activité (préférence pour relations durables)"],
            ["Multisite",    "Forte valeur si présence DROM + Métropole simultanée"],
            ["Stack",        "Migration cuivre → VoIP, remplacement PABX legacy, modernisation"],
        ],
    )

    add_h2(doc, "3.2. Verticales prioritaires")
    add_bullets(doc, [
        "Cabinets d'expertise : avocats, experts-comptables, bureaux d'études (BTP), conseils, assurances",
        "Santé & Médico-social : cliniques, cabinets médicaux, maisons de santé, EHPAD",
        "Commerce & Distribution : enseignes locales, bijouteries (ex. Eurogold/Titeca BEAUPORT), franchises",
        "Hôtellerie & Restauration",
        "Administrations & collectivités",
        "Sociétés d'infogérance (revente B2B2B)",
        "Industrie & BTP",
    ])

    add_h2(doc, "3.3. Persona décideur")
    add_kv_block(doc, [
        ("Rôle",                   "Dirigeant TPE/PME, DSI, responsable IT, office manager"),
        ("Pain points principaux", "Saturation système, coûts opaques, obsolescence (fin du cuivre), manque de centralisation multisite, exigences RGPD"),
        ("Critères de décision",   "Réactivité support · Transparence tarifaire · Certifications · Portabilité numéros · Flexibilité financière"),
    ])

    # === 4. CATALOGUE OFFRES ===
    add_h1(doc, "4. Catalogue des Offres")

    add_h2(doc, "4.1. Trunk SIP (cœur opérateur)")
    add_bullets(doc, [
        "4 / 8 / 16 / 24 / 32 / 64+ appels simultanés",
        "Modes : au compteur (recommandé) ou illimité (sur étude)",
        "Compatibilité : 3CX, Yeastar, Grandstream, Avaya, Alcatel, FreePBX",
        "Sécurité : authentification user/password ou restriction IP",
        "Engagement : 24 mois (E2I) / 36 mois (partenariat SFR Business)",
        "Création de numéros locaux DOM + portabilité entrante",
    ])

    add_h2(doc, "4.2. Solutions IPBX 3CX")
    add_table(doc,
        headers=["Offre", "Cible", "Caractéristiques"],
        rows=[
            ["3CX SMB Pro (mutualisé)",
             "TPE 3-15 utilisateurs",
             "Multisociétés, package clé en main, tarif accessible. Limitations volontaires : pas de CRM/M365, pas de musique perso."],
            ["3CX Pro / Entreprise (dédié)",
             "PME 8+ utilisateurs, multisite",
             "Instance isolée AWS/Azure France-UE, intégrations CRM, multi-tenant, ≥ 90 utilisateurs simultanés possibles."],
        ],
    )
    add_para(doc, "Certification : Partenaire 3CX Silver, expertise 3CX CFD (Call Flow Designer).", italic=True, color=COLOR_SECONDARY)

    add_h2(doc, "4.3. Yeastar / Grandstream (alternative)")
    add_bullets(doc, [
        "Déploiement on-premise ou cloud",
        "Cible : petits cabinets médicaux, structures locales souhaitant garder l'infrastructure en interne",
        "Compatible Trunk SIP E2I ou SFR",
    ])

    add_h2(doc, "4.4. Studio d'enregistrement vocal pro")
    add_bullets(doc, [
        "Messages d'accueil SVI, prédécroché, musique d'attente",
        "1 voix masculine + 2 voix féminines",
        "Livraison sous 48h, formats WAV/MP3/OGG",
        "Optimisé codecs téléphonie (G.711, G.729)",
    ])

    add_h2(doc, "4.5. Assistants vocaux IA (différenciation 2026)")
    add_bullets(doc, [
        "Plateformes : ElevenLabs, Rounded",
        "Cas d'usage : accueil 24/7, qualification leads, prise de RDV, transfert intelligent",
        "Forfaits sans engagement, démo gratuite",
    ])

    add_h2(doc, "4.6. Services complémentaires")
    add_bullets(doc, [
        "Étude de portabilité opérateur",
        "Audit & conseil téléphonie IP",
        "Formations utilisateurs et administrateurs",
        "Déploiement matériel (Fanvil, Yealink, Snom, Patton, Grandstream)",
        "SBC 3CX pour sites multi-postes",
    ])

    # === 5. ATOUTS DIFFÉRENCIANTS ===
    add_h1(doc, "5. Atouts Différenciants")
    add_table(doc,
        headers=["Atout", "Détail"],
        rows=[
            ["Réactivité",
             "Techniciens présents localement (Guyane, Antilles, Métropole). Supervision cloud du parc Fanvil/Yealink, prise en main à distance."],
            ["Expertise",
             "Certifié 3CX/Yeastar, maîtrise CFD & intégrations avancées (HubSpot, Zoho, Salesforce, Organilog, M365, Google Workspace)."],
            ["Hébergement souverain",
             "AWS + Azure France/UE, infrastructure redondée, conformité RGPD, monitoring permanent."],
            ["Proximité",
             "Équipes joignables, offres flexibles, documentation accessible."],
            ["Réseau partenaires",
             "Prestataires IT et infogérants présents dans les DOM pour interventions urgentes sur site."],
            ["Customer Success",
             "Accompagnement structuré : implémentation → formation → suivi continu (revues 30j / 90j / annuelles)."],
        ],
    )

    # === 6. STRATÉGIE TARIFAIRE ===
    add_h1(doc, "6. Stratégie Tarifaire & Modèle Économique")

    add_h2(doc, "6.1. Principes")
    add_bullets(doc, [
        "Récurrence mensuelle dominante (Trunk SIP, IPBX cloud, mobile)",
        "Tarification ajustée à l'usage (au compteur recommandé) ou forfait illimité sur étude",
        "Acompte 50 % sur frais de mise en service et coûts ponctuels",
        "Mensualités démarrent le mois suivant la mise en service",
    ])

    add_h2(doc, "6.2. Modalités de paiement")
    add_bullets(doc, [
        "Virement SEPA",
        "Prélèvement bancaire via GoCardless",
        "Signature électronique via plateforme intégrée HubSpot",
    ])

    add_h2(doc, "6.3. Engagements contractuels")
    add_table(doc,
        headers=["Service", "Engagement"],
        rows=[
            ["Trunk SIP E2I VoIP",                       "24 mois"],
            ["Partenariat SFR Business (Trunk + fibre)", "36 mois"],
            ["Assistants vocaux IA",                     "Sans engagement"],
        ],
    )

    add_h2(doc, "6.4. Financement matériel")
    add_bullets(doc, [
        "Achat direct",
        "Location simple (E2I)",
        "Location financière via Grenke et Cleodis : loyers déductibles, renouvellement/restitution/prolongation, démarche RSE",
    ])

    # === 7. PROCESSUS DE VENTE ===
    add_h1(doc, "7. Processus de Vente")

    add_h2(doc, "7.1. Funnel téléphonie fixe / IPBX 3CX")

    funnel_steps = [
        ("Étape 1 — Capture du lead",
         "Sources : téléphone, formulaire web, email, recommandation, apporteur d'affaires. "
         "Saisie immédiate dans HubSpot CRM avec coordonnées complètes (mail, fixe, mobile, "
         "entreprise) + contexte dans le champ Notes. Renseigner périmètre décisionnel et particularités."),
        ("Étape 2 — Qualification",
         "Premier contact téléphone ou Teams. Comprendre les pain points (saturation, coût, "
         "obsolescence). Adapter le discours en conséquence. Closing : caler une démo produit "
         "Teams immédiatement. Demander la complétion du formulaire Tally."),
        ("Étape 3 — Démonstration produit",
         "Démo 3CX sur Teams, structurée autour des besoins identifiés. Montrer comment la "
         "solution résout les problèmes (pas une simple liste de fonctionnalités). Closing : "
         "caler immédiatement le RDV de présentation du devis. Si Tally non rempli, poster le "
         "lien dans le chat Teams."),
        ("Étape 4 — Présentation du devis",
         "Toujours en visioconférence ou en présentiel (jamais d'envoi mail seul). Expliquer "
         "chaque ligne en lien avec les enjeux du client. Anticiper les objections, dissiper "
         "les doutes, renforcer la perception de valeur."),
        ("Étape 5 — Signature & administratif",
         "Signature électronique HubSpot (recommandée) ou papier + scan. Dossier complet : "
         "formulaire d'ouverture de compte, KBIS/avis INSEE < 3 mois, RIB officiel, CGS "
         "signées, mandat de portabilité + dernière facture opérateur. Email type avec PJ "
         "pré-remplies + rappel téléphonique 48h après envoi. Archivage OneDrive/SharePoint."),
        ("Étape 6 — Activation & Customer Success",
         "Transmission ADV pour mise en service. Suivi à 30 jours (transition, satisfaction). "
         "Suivi à 90 jours (besoins complémentaires, formations, extensions). Revue annuelle "
         "stratégique : positionnement E2I comme partenaire, pas fournisseur."),
    ]
    for title, body in funnel_steps:
        add_h3(doc, title)
        add_para(doc, body)

    add_h2(doc, "7.2. Funnel Trunk SIP (court)")
    add_bullets(doc, [
        "Premier contact : formulaire web ou appel direct (0594 96 35 00)",
        "Recommandation : Trunk au compteur (par défaut) ou illimité (sur étude)",
        "Devis personnalisé après collecte d'infos détaillées",
        "Mise en service + assistance technique de configuration",
        "Compatibilité immédiate : 3CX, Yeastar, Avaya, Alcatel",
        "Tarification au TJM ingénieur pour Asterisk / FreePBX (config plus complexe)",
    ])

    add_h2(doc, "7.3. Leviers psychologiques activés")
    add_table(doc,
        headers=["Levier", "Point d'application"],
        rows=[
            ["Réactivité",            "Saisie CRM immédiate, rappel 48h"],
            ["Réciprocité",           "Suivi 30j/90j sans demande, démo gratuite"],
            ["Preuve sociale",        "Témoignages, badges 3CX, références (Titeca/Eurogold)"],
            ["Engagement progressif", "Tally → démo → devis → signature (escalade par micro-engagements)"],
            ["Cohérence",             "RDV de devis pris dès la fin de la démo"],
        ],
    )

    # === 8. STRATÉGIE MARKETING ===
    add_h1(doc, "8. Stratégie Marketing & Acquisition")

    add_h2(doc, "8.1. Canal principal — Site web e2i-voip.com")
    add_para(doc, "Statut actuel : < 20 leads/mois (sous-performant).", bold=True)
    add_quote(doc,
        "Hypothèse stratégique (révisée après cross-model challenge) : le SEO est un outil de "
        "crédibilité et de closing, pas le moteur d'acquisition principal. Le volume de "
        "recherche local B2B est trop faible pour porter seul l'acquisition. Le site doit "
        "donc convertir les visiteurs déjà ciblés (recommandation, mention, recherche directe)."
    )
    add_para(doc, "Conséquences :", bold=True)
    add_bullets(doc, [
        "Investir dans la conversion (CTA, social proof, clarté du message)",
        "Investir dans le tracking (PostHog) pour mesurer ce qui marche",
        "Reporter à v2 les efforts d'acquisition payants (Google Ads, landing pages)",
    ])

    add_h2(doc, "8.2. Plan d'activation digitale (3 phases GSD)")

    add_h3(doc, "Phase 1 — PostHog Analytics Baseline ✅")
    add_para(doc, "Objectif : tracer anonymement chaque visiteur et capturer les events de conversion.")
    add_bullets(doc, [
        "PostHog JS SDK installé (EU Cloud, GDPR-compliant, persistence: 'memory' = cookieless, pas de cookie banner requis)",
        "3 events instrumentés : cta_click, phone_click, form_submit",
        "Session recordings activés",
        "Try/catch wrapper pour graceful degradation (adblockers)",
    ])

    add_h3(doc, "Phase 2 — Redesign + Social Proof + Conversion ✅")
    add_para(doc, "Objectif : site professionnel avec social proof crédible et chemin de conversion clair.")
    add_bullets(doc, [
        "Social proof : témoignages clients réels + badge 3CX Silver Partner sur homepage et pages produits clés",
        "CTA primaire above-the-fold sur chaque page produit, wording et destination personnalisés",
        "Numéros de téléphone cliquables (tel:) par territoire (Guadeloupe, Martinique, Guyane, La Réunion, France)",
        "Élimination des dead-ends : ContactSectionSimple en bas de chaque page produit",
        "Audit Design System Monolithe 2026 : suppression des violations (font-bold/semibold → font-black uniquement, refactor CTAButton en Link-only)",
    ])

    add_h3(doc, "Phase 3 — Deploy Redesign + SEO Content (en cours)")
    add_para(doc, "Objectif : redesign Monolithe 2026 et contenu SEO en production avec données avant/après dans PostHog.")
    add_bullets(doc, [
        "Merge PR2 (Redesign + Social Proof + Conversion) après période de baseline (1-3 semaines)",
        "Merge PR3 (SEO Content) indépendamment",
        "JSON-LD Service schema sur 7 pages produits (validé)",
        "Sitemap corrigé (14 pages réelles)",
    ])

    add_h2(doc, "8.3. Canaux secondaires")
    add_table(doc,
        headers=["Canal", "Statut", "Priorité"],
        rows=[
            ["Recommandation client",                    "Actif (CSM nourrit la rétention)",          "HAUTE"],
            ["Apporteurs d'affaires / Partenariats",     "Actif (SFR Business, infogérants DOM)",     "HAUTE"],
            ["SEO local",                                "En déploiement (Phase 3)",                  "MOYENNE"],
            ["Tally popup conversion",                   "Actif (déclenchement 3s, ID mDY1bl)",       "MOYENNE"],
            ["Livechat (chat-preoverlay.tsx)",           "Actif",                                     "MOYENNE"],
            ["HubSpot forms",                            "Actif",                                     "HAUTE"],
            ["Google Business Profile",                  "Hors scope v1",                             "DIFFÉRÉ"],
            ["Google Ads",                               "Hors scope v1 (test manuel possible)",      "DIFFÉRÉ"],
            ["LinkedIn / Outbound",                      "Hors périmètre site web",                   "DIFFÉRÉ"],
            ["Blog technique VoIP",                      "Phase ultérieure",                          "DIFFÉRÉ"],
        ],
    )

    add_h2(doc, "8.4. Partenariats commerciaux structurants")
    add_table(doc,
        headers=["Partenaire", "Type", "Apport business"],
        rows=[
            ["3CX (Silver Partner)",     "Vendor",                "Légitimité technique, support 3CX France direct, accès leads écosystème"],
            ["Yeastar",                  "Vendor (depuis 2025)",  "Alternative économique pour petites structures"],
            ["Fanvil / Yealink",         "Hardware",              "Console cloud FDMCS, supervision parc à distance"],
            ["SFR Business Caraïbes",    "Opérateur national",    "Co-commercialisation Fibre + Trunk SIP DOM, légitimation par grand opérateur"],
            ["Grenke / Cleodis",         "Financement",           "Levée d'objection trésorerie chez le prospect"],
        ],
    )

    # === 9. MESSAGING ===
    add_h1(doc, "9. Messaging & Tonalité")

    add_h2(doc, "9.1. Piliers de messaging")
    add_bullets(doc, [
        "Du problème à la solution — clients face à l'arrêt du cuivre, pression budgétaire, complexité multisite",
        "Excellence technique + proximité géographique + tarif transparent — notre triple promesse",
        "Adoption rapide, ROI maximisé, partenariat (pas relation fournisseur) — le résultat client",
    ])

    add_h2(doc, "9.2. Tonalité")
    add_bullets(doc, [
        "Professionnel et rassurant (cible PME averse au risque)",
        "Confiance technique sans jargon excessif",
        "Local et présent — emphase sur les territoires et la réactivité",
        "Concret — exemples chiffrés (« 60 à 90+ utilisateurs », « livraison 48h »)",
    ])

    add_h2(doc, "9.3. Vocabulaire à privilégier")
    add_para(doc,
        "Opérateur de service télécom · intégrateur VoIP · Customer Success · proximité · "
        "expertise certifiée · souverain · transparent · partenaire",
        color=COLOR_MARINE, italic=True
    )

    add_h2(doc, "9.4. Vocabulaire à éviter")
    add_para(doc,
        "« Solution révolutionnaire » · « Leader incontesté » · jargon marketing creux · "
        "promesses non quantifiables",
        color=COLOR_SECONDARY, italic=True
    )

    # === 10. KPIs ===
    add_h1(doc, "10. KPIs & Mesure de Performance")

    add_h2(doc, "10.1. KPIs primaires (Phase Ship & Measure)")
    add_table(doc,
        headers=["KPI", "Source", "Baseline", "Cible 3 mois"],
        rows=[
            ["Leads qualifiés / mois",
             "HubSpot CRM", "< 20", "35-40"],
            ["Taux de conversion cta_click → form_submit",
             "PostHog", "À mesurer", "+30 % vs baseline"],
            ["Taux de clics sur numéros de téléphone",
             "PostHog phone_click", "0 (avant Phase 2)", "À établir"],
            ["Taux d'abandon avant CTA",
             "PostHog session recordings", "À mesurer", "-15 % vs baseline"],
            ["Délai 1er contact lead",
             "HubSpot", "Variable", "< 4h ouvrées"],
        ],
    )

    add_h2(doc, "10.2. KPIs secondaires (suivi commercial)")
    add_bullets(doc, [
        "Taux de conversion qualification → démo (HubSpot)",
        "Taux de conversion démo → devis (HubSpot)",
        "Taux de signature devis → contrat (HubSpot)",
        "Cycle de vente moyen en jours (HubSpot)",
        "Panier moyen (€/mois récurrent) — facturation",
        "Churn rate annuel — facturation",
        "Net Revenue Retention (upsell CSM) — facturation",
        "NPS post-90 jours (enquête CSM)",
    ])

    add_h2(doc, "10.3. Stack analytics")
    add_bullets(doc, [
        "PostHog EU Cloud (cookieless, GDPR-compliant) — comportement site",
        "HubSpot — CRM + marketing automation",
        "Tally — formulaires de qualification approfondie",
        "OneDrive / SharePoint — archivage administratif commercial",
    ])

    # === 11. ENGAGEMENTS ===
    add_h1(doc, "11. Engagements & Garanties")
    add_bullets(doc, [
        "Données hébergées France / UE (AWS, Azure, SFR Cloud)",
        "Conformité RGPD stricte",
        "Services supervisés 24/7",
        "Infrastructure redondée (multi-AZ)",
        "Support technique interne (pas de sous-traitance)",
        "Documentation et formations vidéo accessibles",
        "Customer Success Manager dédié sur la durée du contrat",
    ])

    # === 12. ROADMAP ===
    add_h1(doc, "12. Roadmap Stratégique")

    add_h2(doc, "v1 — en cours (mai/juin 2026)")
    add_bullets(doc, [
        "✅ Tracking comportemental PostHog",
        "✅ Social proof homepage + pages produits clés",
        "✅ Conversion optimisée (CTA, tel:, dead-ends éliminés)",
        "🔄 Déploiement production redesign Monolithe 2026",
        "🔄 Contenu SEO enrichi en production",
        "🔄 Collecte baseline de données conversion",
    ])

    add_h2(doc, "v2 — T3-T4 2026")
    add_bullets(doc, [
        "A/B testing PostHog feature flags (après 1 mois de baseline)",
        "Landing pages dédiées par campagne marketing",
        "Google Business Profile + avis Google intégrés",
        "Cas clients détaillés avec métriques avant/après",
        "Blog technique VoIP (acquisition organique long terme)",
    ])

    add_h2(doc, "v3 — 2027")
    add_bullets(doc, [
        "Outbound LinkedIn / prospection structurée",
        "Programme apporteurs d'affaires formalisé",
        "Marketplace assistants vocaux IA (différenciation marché)",
        "Programme de certification revendeurs régionaux",
    ])

    # === 13. HORS SCOPE ===
    add_h1(doc, "13. Hors Scope (explicite)")
    add_table(doc,
        headers=["Élément", "Raison"],
        rows=[
            ["Refonte complète du contenu éditorial",  "Déjà fait sur la branche dev"],
            ["Backend / API custom",                    "Site statique sur Vercel, pas de besoin"],
            ["Mobile app dédiée",                       "Site responsive suffisant"],
            ["Google Ads campagnes",                    "Optionnel, budget minimal, à tester manuellement"],
            ["Outbound LinkedIn / prospection",         "Hors périmètre site web (canal séparé v3)"],
        ],
    )

    # === 14. ANNEXES ===
    add_h1(doc, "14. Annexes")

    add_h2(doc, "14.1. Numéros commerciaux par territoire")
    add_table(doc,
        headers=["Territoire", "Numéro"],
        rows=[
            ["Guyane",                  "05 94 96 35 00"],
            ["Guadeloupe",              "05 90 17 35 00"],
            ["Martinique",              "05 96 31 35 00"],
            ["La Réunion",              "+262 263 08 55 00"],
            ["France métropolitaine",   "01 89 56 35 00"],
            ["Hotline support technique", "01 89 56 05 00"],
        ],
    )

    add_h2(doc, "14.2. Contacts officiels")
    add_kv_block(doc, [
        ("Site web",                  "https://www.e2i-voip.com"),
        ("Email commercial",          "commerciaux@e2i-voip.com"),
        ("Email support",             "assistance@e2i-voip.com"),
        ("Demande de devis en ligne", "https://www.e2i-voip.com/devis"),
        ("Portail support",           "https://support.e2i-voip.com"),
    ])

    add_h2(doc, "14.3. Charte graphique appliquée à ce document")
    add_table(doc,
        headers=["Rôle", "Hex", "Usage dans le document"],
        rows=[
            ["Rouge Principal",   "#E53E3E", "Titres H1, séparateurs, accents, citations"],
            ["Bleu Marine",       "#2D3848", "Titres H2/H3, en-têtes de tableaux"],
            ["Gris Foncé",        "#1F2937", "Texte courant"],
            ["Gris Secondaire",   "#818096", "Textes secondaires, légendes, vocabulaire à éviter"],
        ],
    )

    # === FOOTER ===
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    add_horizontal_rule(p, color_rgb=COLOR_MARINE, thickness_pt=1)
    r = p.add_run(
        "Document de spécification — E2I VoIP — Stratégie de Vente & Marketing v1.0 — 2026-04-28"
    )
    style_run(r, color=COLOR_SECONDARY, italic=True, size=9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    out = Path(__file__).resolve().parent.parent / "docs" / "SPEC_STRATEGIE_VENTE_MARKETING.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    saved = build_document(str(out))
    print(f"OK · Generated: {saved}")
